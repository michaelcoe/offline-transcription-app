from pathlib import Path
# importing modules 
from docx import Document

import numpy as np

transcript_file = Path('./audio/')

temp_transcript = str(transcript_file.name) + '.docx'
document = Document()
p = document.add_paragraph()

with open(transcript_file, 'r') as f:
    lines = f.readlines()

for line in lines:
    p.add_run(line)

document.save(temp_transcript)

# Remove the original transcript and then replace with newly created file
transcript_file = Path(temp_transcript)
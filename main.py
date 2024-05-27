import streamlit as st
import numpy as np
import codecs
import tempfile
import os 
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import A4
from reportlab.platypus import Paragraph
from reportlab.platypus import SimpleDocTemplate
from pathlib import Path
import subprocess
from docx import Document

st.set_page_config(
    page_title="Offline Transcription",
    layout="centered",
    initial_sidebar_state="auto"
)

if 'transcript' not in st.session_state:
    st.session_state['transcript'] = None

if 'transcript_file' not in st.session_state:
    temp_transcript = Path('./audio/transcript.srt')
    with open(temp_transcript, 'w') as tt:
        tt.write(" ")
    
    st.session_state['transcript_file'] = Path('./audio/transcript.srt')
    st.session_state['transcript_output'] = st.session_state['transcript_file']

#@st.cache_resource(show_spinner="Transcribing...")
def transcription(audio_file, model):
    if st.session_state['output'] == 'pdf' or st.session_state['output'] == 'docx':
        temp_output = 'srt'
    else:
        temp_output = st.session_state['output']

    if st.session_state['eo'] == 'yes':
        cmd = f'./Whisper-Faster-XXL/whisper-faster-xxl ./{str(audio_file)} --language English --model {model} --output_format {temp_output} --output_dir source'.split()
    else:
        cmd = f'./Whisper-Faster-XXL/whisper-faster-xxl ./{str(audio_file)} --model {model} --output_format {temp_output} --output_dir source'.split()      
    
    # try the transcription or throw an error
    #return_code = subprocess.run(cmd, shell=False)
    return_code = subprocess.run(cmd, shell=False, stdout=subprocess.DEVNULL, stderr=subprocess.DEVNULL)

    st.session_state['transcript_file'] = Path(str(audio_file.parent.joinpath(audio_file.stem)) + '.' + temp_output)

    if return_code.check_returncode() == None:
        with codecs.open(st.session_state['transcript_file'], encoding='utf-8') as file:
            data = file.read()

    st.session_state['transcript'] = data
    st.session_state['transcript_output'] = st.session_state['transcript_file']

    return return_code.check_returncode()

def convert_transcript():
    export = st.session_state['export']
    transcript_file = st.session_state['transcript_file']
    new_transcript = Path(str(transcript_file.parent.joinpath(transcript_file.stem)) + '.' + st.session_state['export'])
    
    st.session_state['transcript_output'] = new_transcript
    
    if export == 'pdf':
        text_width=A4[0] / 2
        text_height=A4[1] / 2

        # creating a pdf object 
        pdf = SimpleDocTemplate(str(new_transcript))
        styles = getSampleStyleSheet()

        with open(st.session_state['transcript_file'], 'r') as file:
            lines = file.readlines()

        paragraphs = []

        for i in np.arange(0, len(lines), 3):
            line = ''.join(lines[i:i+3]).replace('\n','<br/>\n')
            # line = line.replace('\n','<br/>\n')

            p = Paragraph(line, styles["Normal"])   
            p.wrapOn(pdf, text_width, text_height)
            paragraphs.append(p)

        pdf.build(paragraphs)

    elif export == 'docx':
        document = Document()
        p = document.add_paragraph()
        
        with open(st.session_state['transcript_file'], 'r') as f:
            lines = f.readlines()
        
        for line in lines:
            p.add_run(line)
        
        document.save(new_transcript)    
    else:
        with open(st.session_state['transcript_output'], 'w') as tr:
            tr.write(st.session_state['transcript'])

if __name__ == "__main__":
    # ------------------- Sidebar Information -------------------------
    st.title('Offline Transcription App')

    # UC banner
    if Path('./UCWhite.png').is_file():
        st.sidebar.image('UCWhite.png')
    else:
        pass
    # Add a description in the sidebar
    st.sidebar.title('About this app')
    st.sidebar.markdown("""This app uses the offline version of the openAI Whisper Automatic Speech Recognition (ASR) package to transcribe uploaded audio or video files. 
                    To transcribe an audio or video file, drag and drop the file, or you can use the **Browse Files** button""")

    st.sidebar.subheader("Note")
    st.sidebar.markdown("""This Whisper ASR package has been trained using machine learning, but is secure to use at UC. Your data will not be used to train future versions 
                        of the app. All files are automatically deleted after closing the browser window. Please ensure that you download the generated transcript file and
                        save it in a secure location. You should also save the original audio or video file in a secure location.
                        """)

    st.sidebar.subheader("Support")
    st.sidebar.markdown("""Our eResearch consultants are on hand to support your use of this app and for support with data storage. For support, please contact the eResearch
                        team using UC services [eResearch consultancy form](https://services.canterbury.ac.nz/uc?id=sc_cat_item&sys_id=8effe377db992510e447f561f396197c)""")

    # ------------------------ Audio File form ---------------------------

    # Model and audio file in a form
    with st.form("setup-form", clear_on_submit=True):
        model_select = st.radio('Select a model', 
                        ['tiny', 'base', 'small', 'medium', 'large', 'large-v2', 'large-v3'],
                        key='model',
                        index=3,
                        horizontal=True)
        eo = st.radio('English Only', 
                        ['yes', 'no'],
                        key='eo',
                        index=1,
                        horizontal=True)
        
        # File uploader
        uploaded_file = st.file_uploader(
            "Upload file you want to transcribe",
            accept_multiple_files=False,
        )
        
        output_select = st.radio('Select an output format',
                                ['srt', 'json'],
                                key='output',
                                index=0,
                                horizontal=True)
        
        transcribe_btn = st.form_submit_button("Transcribe")

    if transcribe_btn == True: 
        if uploaded_file is not None:
            # write the audio file to a folder
            audio_file=Path('./audio').joinpath(uploaded_file.name)        
            with open(audio_file, 'wb') as f:
                f.write(uploaded_file.getbuffer())
            
                # Transcribe the audio file
                if st.session_state['eo'] == 'yes':
                    model = st.session_state['model'] + '.en'
                else:
                    model = st.session_state['model']
                
                with st.spinner("Transcribing!..."):
                    files = list(Path('./audio').rglob("*.srt"))
                    for file in files:
                        os.remove(file)
                    return_code = transcription(audio_file, model)

                if return_code == None:
                    st.success('Transcription complete!')
                    st.session_state['disabled'] = False              
                    os.remove(audio_file)
                else:
                    st.warning("Something went wrong. Please contact the eResearch Team. Or try refreshing the app.")

                with st.expander(label='Preview the transcript'):
                    st.write(st.session_state['transcript'])
                # Download the transcript
        
    # Output widgets
    col1, col2 = st.columns(2)

    if 'export' not in st.session_state:
        st.session_state['export'] = 'srt'
        st.session_state['disabled'] = True

    with col1:
        if st.session_state['output'] == 'srt':
            export_select = st.selectbox('Select your file export format',
                                ['srt', 'txt', 'docx', 'pdf', 'vtt', 'lrc', 'tsv'],
                                key='export',
                                index=0,
                                on_change=convert_transcript,
                                disabled=st.session_state['disabled'],
                                )
        else:
            export_select = st.selectbox('Select your file export format',
                    ['json'],
                    key='export',
                    index=0,
                    disabled=st.session_state['disabled'],
                    )
    
        if st.session_state['export'] == 'docx':
            mime = 'docx'
        elif st.session_state['export'] == 'json':
            mime = 'application/json'
        elif st.session_state['export'] == 'pdf':
            mime = 'application/octet-stream'
        else:
            mime = 'text/plain'    
    
    with col2:
        st.write(" ")
        st.write(" ")
        with open(st.session_state['transcript_output'], 'rb') as f: 
            download = st.download_button(
                        label='Download Transcript',
                        data = f,
                        file_name=st.session_state['transcript_output'].name,
                        mime = mime,
                        disabled=st.session_state['disabled'],
                        )
    
            if download:
                files = list(Path('./audio').rglob("*.*"))
                for file in files:
                    if not file.suffix == ".srt":
                        os.remove(file)

    #st.write(st.session_state)
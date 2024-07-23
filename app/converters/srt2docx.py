from pathlib import Path
from docx import Document

def convert(transcript_file, transcript):
    new_transcript = Path(str(transcript_file.parent.joinpath(transcript_file.stem)) + '.docx')

    document = Document()
    p = document.add_paragraph()
        
    lines = transcript.split('\n')
    # need to add carrage return for proper formatting
    for i, line in enumerate(lines):
        lines[i] = lines[i] + '\n'
        
    for line in lines:
        p.add_run(line)
        
    document.save(new_transcript)

    return new_transcript
